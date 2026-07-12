"""
DOCX Parser

Extracts text content from DOCX files.
"""

import logging
from typing import Dict, Any, Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..processor import ProcessedFile, FileType


logger = logging.getLogger(__name__)


class DocxParser:
    """Parser for DOCX files."""

    def __init__(self):
        if not DOCX_AVAILABLE:
            logger.warning("[DOCX_PARSER] python-docx not installed")

    async def parse(
        self,
        file_data: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedFile:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_data: Raw DOCX bytes
            filename: Original filename
            metadata: Additional metadata
            
        Returns:
            ProcessedFile with extracted text
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            import io
            
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n".join(text_content)
            
            # Extract metadata
            docx_metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            if doc.core_properties:
                docx_metadata.update({
                    "title": doc.core_properties.title or '',
                    "author": doc.core_properties.author or '',
                    "subject": doc.core_properties.subject or '',
                    "created": str(doc.core_properties.created) if doc.core_properties.created else '',
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ''
                })
            
            # Merge with provided metadata
            if metadata:
                docx_metadata.update(metadata)
            
            return ProcessedFile(
                filename=filename,
                file_type=FileType.DOCX,
                content=full_text,
                metadata=docx_metadata
            )
            
        except Exception as e:
            logger.error(f"[DOCX_PARSER] Failed to parse DOCX: {e}")
            raise
